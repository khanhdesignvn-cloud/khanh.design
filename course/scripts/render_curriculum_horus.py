#!/usr/bin/env python3
"""Dựng giáo án giảng viên chi tiết theo visual system Horus."""
from __future__ import annotations

import argparse
import re
from html import escape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents

ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "course" / "curriculum"
ASSETS = ROOT / "course" / "assets"
DIST = ROOT / "course" / "dist"
PREFIX = "04-09-2026"
DEFAULT_PDF = DIST / f"{PREFIX}-Giao-an-chi-tiet-Horus-AI-Van-hanh-doanh-nghiep.pdf"
DEFAULT_DOCX = DIST / f"{PREFIX}-Giao-an-chi-tiet-Horus-AI-Van-hanh-doanh-nghiep.docx"

FONT_FAMILY = "Google Sans Flex 24pt"
FONT_REG = ASSETS / "GoogleSansFlex-400.ttf"
FONT_MED = ASSETS / "GoogleSansFlex-500.ttf"
FONT_BOLD = ASSETS / "GoogleSansFlex-700.ttf"

INK = colors.HexColor("#161616")
WHITE = colors.white
MUTED = colors.HexColor("#858585")
LINE = colors.HexColor("#D4D4D4")
SOFT = colors.HexColor("#F5F5F5")
ALERT = colors.HexColor("#FFF4E5")
GREEN = colors.HexColor("#EAF8EF")

WEEK_META = {
    1: ("LÀM CHỦ CLAUDE & TƯ DUY GIAO VIỆC", "Bản đồ cơ hội AI và ba quy trình ưu tiên", "Chọn việc đúng · Viết brief · Ẩn danh dữ liệu"),
    2: ("XÂY TRỢ LÝ HIỂU DOANH NGHIỆP", "Claude Project, hồ sơ doanh nghiệp và checklist chất lượng", "Nguồn dữ liệu · Xung đột · Giọng thương hiệu"),
    3: ("XÂY HỆ THỐNG NỘI DUNG MARKETING", "Lịch nội dung 7 ngày và gói Facebook–email–video", "Nguồn xác minh · Đa kênh · Duyệt xuất bản"),
    4: ("XÂY TRỢ LÝ BÁN HÀNG", "Kịch bản khám phá, xử lý từ chối và chuỗi theo dõi ba chạm", "Tư vấn · Điều kiện dừng · Không tự quyết giá"),
    5: ("CHĂM SÓC KHÁCH HÀNG & VẬN HÀNH", "Bộ FAQ–chuyển cấp và một SOP nội bộ", "Phân loại · Chuyển cấp · Bằng chứng hoàn tất"),
    6: ("TỰ ĐỘNG HÓA & BÀN GIAO HỆ THỐNG", "Workflow chạy thử và hồ sơ bàn giao bộ máy AI", "Validation · Chống trùng · Runbook"),
}


def register_fonts():
    for path in (FONT_REG, FONT_MED, FONT_BOLD):
        if not path.exists():
            raise FileNotFoundError(path)
    pdfmetrics.registerFont(TTFont("GSF", str(FONT_REG)))
    pdfmetrics.registerFont(TTFont("GSF-Medium", str(FONT_MED)))
    pdfmetrics.registerFont(TTFont("GSF-Bold", str(FONT_BOLD)))
    pdfmetrics.registerFontFamily("GSF", normal="GSF", bold="GSF-Bold")


def inline(text: str) -> str:
    value = escape(text.strip())
    value = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", value)
    value = re.sub(r"`(.+?)`", r"<font name='GSF'>\1</font>", value)
    return value


def pdf_styles():
    getSampleStyleSheet()
    return {
        "H1": ParagraphStyle("H1", fontName="GSF-Bold", fontSize=18, leading=22, textColor=INK, spaceBefore=6*mm, spaceAfter=3*mm, keepWithNext=True),
        "H2": ParagraphStyle("H2", fontName="GSF-Bold", fontSize=13, leading=17, textColor=INK, spaceBefore=5*mm, spaceAfter=2*mm, keepWithNext=True, borderColor=INK, borderWidth=0, borderPadding=0),
        "H3": ParagraphStyle("H3", fontName="GSF-Bold", fontSize=10.2, leading=14, textColor=INK, spaceBefore=3.5*mm, spaceAfter=1.5*mm, keepWithNext=True),
        "Body": ParagraphStyle("Body", fontName="GSF", fontSize=8.8, leading=13.2, textColor=INK, alignment=TA_LEFT, spaceAfter=2.1*mm),
        "Bullet": ParagraphStyle("Bullet", fontName="GSF", fontSize=8.8, leading=13.2, leftIndent=5*mm, firstLineIndent=-3.3*mm, textColor=INK, spaceAfter=1.3*mm),
        "Note": ParagraphStyle("Note", fontName="GSF", fontSize=8.1, leading=12, textColor=MUTED, spaceAfter=1.5*mm),
        "Callout": ParagraphStyle("Callout", fontName="GSF", fontSize=8.8, leading=13.2, textColor=INK, leftIndent=3*mm, rightIndent=3*mm, borderColor=LINE, borderWidth=.6, borderPadding=6, backColor=SOFT, spaceAfter=3*mm),
        "Table": ParagraphStyle("Table", fontName="GSF", fontSize=7.1, leading=9.8, textColor=INK),
        "TableHead": ParagraphStyle("TableHead", fontName="GSF-Bold", fontSize=7.1, leading=9.5, textColor=WHITE),
        "TOC0": ParagraphStyle("TOC0", fontName="GSF-Bold", fontSize=10.5, leading=15, textColor=INK, leftIndent=0, spaceBefore=2*mm),
        "TOC1": ParagraphStyle("TOC1", fontName="GSF", fontSize=8.5, leading=12.5, textColor=MUTED, leftIndent=7*mm),
    }


class HeroPanel(Flowable):
    def __init__(self, overline: str, title: str, subtitle: str, meta_left: str, meta_right: str, height=72*mm, toc_level=None):
        super().__init__()
        self.overline, self.title, self.subtitle = overline, title, subtitle
        self.meta_left, self.meta_right = meta_left, meta_right
        self.height = height
        self.toc_title = title
        self.toc_level = toc_level

    def wrap(self, avail_width, avail_height):
        self.width = avail_width
        return avail_width, self.height

    def draw(self):
        c = self.canv
        c.setFillColor(INK)
        c.roundRect(0, 0, self.width, self.height, 8*mm, fill=1, stroke=0)
        c.setFillColor(colors.HexColor("#C9C9C9")); c.setFont("GSF-Medium", 7.2)
        c.drawString(10*mm, self.height-12*mm, self.overline)
        c.setFillColor(WHITE); c.setFont("GSF-Medium", 21)
        y = self.height-27*mm
        for line in self.title.split("\n"):
            c.drawString(10*mm, y, line); y -= 8.2*mm
        c.setFillColor(colors.HexColor("#B9B9B9")); c.setFont("GSF", 9)
        c.drawString(10*mm, y+1*mm, self.subtitle)
        c.setStrokeColor(colors.HexColor("#404040")); c.line(10*mm, 17*mm, self.width-10*mm, 17*mm)
        c.setFont("GSF", 6.7); c.setFillColor(colors.HexColor("#929292"))
        c.drawString(10*mm, 10.5*mm, self.meta_left.upper())
        c.drawRightString(self.width-10*mm, 10.5*mm, self.meta_right.upper())


class PartHeading(Flowable):
    def __init__(self, label: str, title: str):
        super().__init__(); self.label=label; self.title=title; self.height=13*mm
        self.toc_title=title; self.toc_level=1

    def wrap(self, avail_width, avail_height):
        self.width=avail_width; return avail_width,self.height

    def draw(self):
        c=self.canv
        c.setFillColor(INK); c.roundRect(0,2.2*mm,24*mm,7.2*mm,3.6*mm,fill=1,stroke=0)
        c.setFillColor(WHITE); c.setFont("GSF-Medium",6.7); c.drawCentredString(12*mm,4.7*mm,self.label.upper())
        c.setFillColor(INK); c.setFont("GSF-Bold",12); c.drawString(29*mm,4.1*mm,self.title.upper())


def pdf_table(rows, st):
    cols=max(len(r) for r in rows); normalized=[r+[""]*(cols-len(r)) for r in rows]
    data=[]
    for ri,row in enumerate(normalized):
        data.append([Paragraph(inline(x),st["TableHead" if ri==0 else "Table"]) for x in row])
    available=A4[0]-34*mm
    weights=[]
    for ci in range(cols):
        longest=max(len(r[ci]) for r in normalized)
        weights.append(max(1,min(4.8,longest/18)))
    widths=[available*x/sum(weights) for x in weights]
    table=Table(data,colWidths=widths,repeatRows=1,hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),INK),("TEXTCOLOR",(0,0),(-1,0),WHITE),
        ("BACKGROUND",(0,1),(-1,-1),WHITE),("ROWBACKGROUNDS",(0,1),(-1,-1),[WHITE,SOFT]),
        ("BOX",(0,0),(-1,-1),.45,LINE),("INNERGRID",(0,0),(-1,-1),.3,LINE),
        ("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),5),
        ("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    return table


def markdown_pdf(text: str, st, skip_h1=True):
    flows=[]; lines=text.splitlines(); i=0; h2_count=0
    while i<len(lines):
        line=lines[i].strip()
        if not line: i+=1; continue
        if line.startswith("|") and i+1<len(lines) and re.match(r"^\|?[\s:|-]+\|?$",lines[i+1].strip()):
            rows=[]
            while i<len(lines) and lines[i].strip().startswith("|"):
                cells=[c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[\s:|-]+",c or " ") for c in cells): rows.append(cells)
                i+=1
            flows += [pdf_table(rows,st),Spacer(1,2.5*mm)]; continue
        m=re.match(r"^(#{1,3})\s+(.+)$",line)
        if m:
            level=len(m.group(1)); title=m.group(2)
            if level==1 and skip_h1: i+=1; continue
            if level==2:
                h2_count+=1; flows.append(PartHeading(f"MỤC {h2_count:02d}",title))
            else: flows.append(Paragraph(inline(title),st[f"H{level}"]))
            i+=1; continue
        if line.startswith("> "):
            flows.append(Paragraph(inline(line[2:]),st["Callout"])); i+=1; continue
        bullet=re.match(r"^[-*]\s+(.+)$",line); numbered=re.match(r"^(\d+)\.\s+(.+)$",line)
        if bullet or numbered:
            prefix="•" if bullet else numbered.group(1)+"."
            body=bullet.group(1) if bullet else numbered.group(2)
            flows.append(Paragraph(prefix+" "+inline(body),st["Bullet"])); i+=1; continue
        para=[line]; i+=1
        while i<len(lines):
            nxt=lines[i].strip()
            if not nxt or nxt.startswith(("#","|","> ","- ","* ")) or re.match(r"^\d+\.\s+",nxt): break
            para.append(nxt); i+=1
        joined=" ".join(para)
        special=any(x in joined for x in ("Điểm dừng bắt buộc","DỰ KIẾN","Cần sửa"))
        flows.append(Paragraph(inline(joined),st["Callout" if special else "Body"]))
    return flows


class HorusDoc(BaseDocTemplate):
    def __init__(self, output, st):
        super().__init__(str(output),pagesize=A4,leftMargin=17*mm,rightMargin=17*mm,topMargin=14*mm,bottomMargin=16*mm,title="Giáo án chi tiết AI Vận Hành Doanh Nghiệp",author="khanh.design")
        frame=Frame(self.leftMargin,self.bottomMargin,self.width,self.height,id="body")
        self.addPageTemplates(PageTemplate(id="main",frames=frame,onPage=self.draw_page))
        self.seq=0; self.st=st
    def beforeDocument(self): self.seq=0; super().beforeDocument()
    def draw_page(self,c,doc):
        c.saveState(); c.setStrokeColor(LINE); c.line(17*mm,12.5*mm,A4[0]-17*mm,12.5*mm)
        c.setFillColor(MUTED); c.setFont("GSF",6.8); c.drawString(17*mm,7.8*mm,"KHANH.DESIGN — AI VẬN HÀNH DOANH NGHIỆP")
        c.drawRightString(A4[0]-17*mm,7.8*mm,f"{doc.page} | GIÁO ÁN GIẢNG VIÊN"); c.restoreState()
    def afterFlowable(self,f):
        title=None; level=None
        if isinstance(f,(HeroPanel,PartHeading)) and f.toc_level is not None: title=f.toc_title.replace("\n"," "); level=f.toc_level
        elif isinstance(f,Paragraph) and f.style.name in ("H1","H2"): title=f.getPlainText(); level=0 if f.style.name=="H1" else 1
        if title is not None:
            key=f"h{self.seq}"; self.seq+=1; self.canv.bookmarkPage(key); self.canv.addOutlineEntry(title,key,level=level,closed=level==0); self.notify("TOCEntry",(level,title,self.page,key))


def quick_dashboard(week, st):
    title,product,focus=WEEK_META[week]
    data=[
        [Paragraph("ĐẦU RA BẮT BUỘC",st["TableHead"]),Paragraph("NHỊP ĐỨNG LỚP",st["TableHead"])],
        [Paragraph(inline(product),st["Table"]),Paragraph("10' kiểm tra · 20' tư duy · 25' demo · 65' thực hành · 20' nhận xét · 10' giao bài",st["Table"])],
        [Paragraph("TRỌNG TÂM",st["TableHead"]),Paragraph("ĐIỂM DỪNG",st["TableHead"])],
        [Paragraph(inline(focus),st["Table"]),Paragraph("Phút 120: đóng phạm vi, chuyển sang kiểm tra chất lượng và an toàn.",st["Table"])],
    ]
    t=Table(data,colWidths=[(A4[0]-34*mm)/2]*2,hAlign="LEFT")
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),INK),("BACKGROUND",(0,2),(-1,2),INK),("BOX",(0,0),(-1,-1),.5,LINE),("INNERGRID",(0,0),(-1,-1),.35,LINE),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6)]))
    return t


def build_pdf(output: Path):
    register_fonts(); st=pdf_styles(); output.parent.mkdir(parents=True,exist_ok=True)
    toc=TableOfContents(); toc.levelStyles=[st["TOC0"],st["TOC1"]]
    story=[
        Spacer(1,4*mm),HeroPanel("KHANH.DESIGN — HỆ THỐNG ĐÀO TẠO", "GIÁO ÁN GIẢNG VIÊN\nCHI TIẾT 6 TUẦN", "AI VẬN HÀNH DOANH NGHIỆP", "Claude là công cụ trung tâm", "Phiên bản 04/09/2026",82*mm,None),
        Spacer(1,12*mm),Paragraph("TÀI LIỆU DÙNG KHI ĐỨNG LỚP",st["H1"]),Paragraph("Bao gồm hướng dẫn điều phối chung và sáu giáo án 150 phút. Mỗi tuần có bảng điều khiển nhanh, lời dẫn, demo, thực hành, rubric, bài tập, điểm dừng an toàn và phương án xử lý lớp.",st["Body"]),
        PageBreak(),PartHeading("CHỈ MỤC","MỤC LỤC GIÁO ÁN"),Spacer(1,4*mm),toc,PageBreak(),
        HeroPanel("PHẦN 00 — VẬN HÀNH LỚP", "HƯỚNG DẪN\nGIẢNG VIÊN", "Nhịp lớp · Phản hồi · Dữ liệu · Dự phòng", "06 buổi × 150 phút", "06 AI Clinic × 60 phút",58*mm,0),Spacer(1,5*mm)
    ]
    story.extend(markdown_pdf((SOURCE/"facilitator-guide.md").read_text(encoding="utf-8"),st,True))
    for week in range(1,7):
        title,product,focus=WEEK_META[week]
        story += [PageBreak(),HeroPanel(f"PHẦN {week:02d} — TUẦN {week}",title.replace(" & "," &\n") if len(title)>34 else title,product,"Buổi chính 150 phút","AI Clinic 60 phút",60*mm,0),Spacer(1,5*mm),quick_dashboard(week,st),Spacer(1,4*mm)]
        text=(SOURCE/f"week-{week:02d}.md").read_text(encoding="utf-8")
        # Metadata is represented in the hero/dashboard; remove only those three leading lines.
        lines=text.splitlines(); body="\n".join(lines[5:])
        story.extend(markdown_pdf(body,st,True))
    doc=HorusDoc(output,st); doc.multiBuild(story); return output


# ---------- Editable DOCX ----------
def shade(cell,color):
    tcpr=cell._tc.get_or_add_tcPr(); shd=tcpr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcpr.append(shd)
    shd.set(qn("w:fill"),color)


def set_run_font(run,size=9,bold=False,color="161616",weight=None):
    run.font.name=FONT_FAMILY; run._element.rPr.rFonts.set(qn("w:eastAsia"),FONT_FAMILY); run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)


def docx_para(doc,text,style=None,bold=False,color="161616",size=9):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(4); r=p.add_run(text); set_run_font(r,size,bold,color); return p


def docx_hero(doc,overline,title,subtitle,left,right):
    table=doc.add_table(rows=1,cols=1); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    cell=table.cell(0,0); shade(cell,"161616"); cell.margin_top=Cm(.7); cell.margin_bottom=Cm(.7)
    p=cell.paragraphs[0]; p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(8)
    r=p.add_run(overline); set_run_font(r,7,False,"C9C9C9")
    p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(4); r=p.add_run(title); set_run_font(r,22,True,"FFFFFF")
    p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(14); r=p.add_run(subtitle); set_run_font(r,9,False,"B9B9B9")
    p=cell.add_paragraph(); r=p.add_run(f"{left.upper()}                                      {right.upper()}"); set_run_font(r,7,False,"929292")
    doc.add_paragraph()


def docx_table(doc,rows):
    cols=max(len(r) for r in rows); table=doc.add_table(rows=0,cols=cols); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for ci in range(cols):
            text=row[ci] if ci<len(row) else ""; cells[ci].text=""; p=cells[ci].paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(text); set_run_font(r,7.5,ri==0,"FFFFFF" if ri==0 else "161616"); cells[ci].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cells[ci],"161616" if ri==0 else ("F5F5F5" if ri%2==0 else "FFFFFF"))
    doc.add_paragraph(); return table


def markdown_docx(doc,text,skip_h1=True):
    lines=text.splitlines(); i=0; h2=0
    while i<len(lines):
        line=lines[i].strip()
        if not line: i+=1; continue
        if line.startswith("|") and i+1<len(lines) and re.match(r"^\|?[\s:|-]+\|?$",lines[i+1].strip()):
            rows=[]
            while i<len(lines) and lines[i].strip().startswith("|"):
                cells=[re.sub(r"\*\*|`","",c.strip()) for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[\s:|-]+",c or " ") for c in cells): rows.append(cells)
                i+=1
            docx_table(doc,rows); continue
        m=re.match(r"^(#{1,3})\s+(.+)$",line)
        if m:
            level=len(m.group(1)); title=re.sub(r"\*\*|`","",m.group(2))
            if level==1 and skip_h1: i+=1; continue
            if level==2:
                h2+=1; table=doc.add_table(rows=1,cols=2); table.autofit=False; table.columns[0].width=Cm(2.3); table.columns[1].width=Cm(14.4)
                c1,c2=table.rows[0].cells; shade(c1,"161616"); c1.text=""; r=c1.paragraphs[0].add_run(f"MỤC {h2:02d}"); set_run_font(r,7,True,"FFFFFF"); c2.text=""; r=c2.paragraphs[0].add_run(title.upper()); set_run_font(r,12,True,"161616")
                doc.add_paragraph()
            else: docx_para(doc,title,style="Heading 1" if level==1 else "Heading 3",bold=True,color="161616",size=16 if level==1 else 10.5)
            i+=1; continue
        bullet=re.match(r"^[-*]\s+(.+)$",line); numbered=re.match(r"^(\d+)\.\s+(.+)$",line)
        if bullet or numbered:
            body=bullet.group(1) if bullet else numbered.group(2); style="List Bullet" if bullet else "List Number"; docx_para(doc,re.sub(r"\*\*|`","",body),style=style,size=8.8); i+=1; continue
        para=[line]; i+=1
        while i<len(lines):
            nxt=lines[i].strip()
            if not nxt or nxt.startswith(("#","|","> ","- ","* ")) or re.match(r"^\d+\.\s+",nxt): break
            para.append(nxt); i+=1
        docx_para(doc,re.sub(r"\*\*|`",""," ".join(para)),size=8.8)


def build_docx(output: Path):
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Cm(1.4); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(1.7); sec.right_margin=Cm(1.7)
    for style_name in ["Normal","Heading 1","Heading 2","Heading 3","List Bullet","List Number"]:
        st=doc.styles[style_name]; st.font.name=FONT_FAMILY; st._element.rPr.rFonts.set(qn("w:eastAsia"),FONT_FAMILY); st.font.size=Pt(8.8); st.font.color.rgb=RGBColor(22,22,22)
    docx_hero(doc,"KHANH.DESIGN — HỆ THỐNG ĐÀO TẠO","GIÁO ÁN GIẢNG VIÊN\nCHI TIẾT 6 TUẦN","AI VẬN HÀNH DOANH NGHIỆP","Claude là công cụ trung tâm","Phiên bản 04/09/2026")
    docx_para(doc,"TÀI LIỆU DÙNG KHI ĐỨNG LỚP",bold=True,size=16); docx_para(doc,"Hướng dẫn điều phối chung và sáu giáo án 150 phút: lời dẫn, demo, thực hành, rubric, bài tập, điểm dừng và phương án xử lý lớp.",size=10)
    doc.add_page_break(); docx_hero(doc,"PHẦN 00 — VẬN HÀNH LỚP","HƯỚNG DẪN GIẢNG VIÊN","Nhịp lớp · Phản hồi · Dữ liệu · Dự phòng","06 buổi × 150 phút","06 AI Clinic × 60 phút")
    markdown_docx(doc,(SOURCE/"facilitator-guide.md").read_text(encoding="utf-8"),True)
    for week in range(1,7):
        doc.add_page_break(); title,product,focus=WEEK_META[week]; docx_hero(doc,f"PHẦN {week:02d} — TUẦN {week}",title,product,"Buổi chính 150 phút","AI Clinic 60 phút")
        docx_table(doc,[["ĐẦU RA BẮT BUỘC","NHỊP ĐỨNG LỚP"],[product,"10' kiểm tra · 20' tư duy · 25' demo · 65' thực hành · 20' nhận xét · 10' giao bài"],["TRỌNG TÂM","ĐIỂM DỪNG"],[focus,"Phút 120: đóng phạm vi, chuyển sang kiểm tra chất lượng và an toàn."]])
        lines=(SOURCE/f"week-{week:02d}.md").read_text(encoding="utf-8").splitlines(); markdown_docx(doc,"\n".join(lines[5:]),True)
    for section in doc.sections:
        p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("KHANH.DESIGN — AI VẬN HÀNH DOANH NGHIỆP · GIÁO ÁN GIẢNG VIÊN"); set_run_font(r,7,False,"858585")
    output.parent.mkdir(parents=True,exist_ok=True); doc.save(output); return output


def main():
    parser=argparse.ArgumentParser(); parser.add_argument("--pdf",type=Path,default=DEFAULT_PDF); parser.add_argument("--docx",type=Path,default=DEFAULT_DOCX); args=parser.parse_args()
    print(build_pdf(args.pdf)); print(build_docx(args.docx))


if __name__=="__main__": main()
