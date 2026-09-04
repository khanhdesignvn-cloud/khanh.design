#!/usr/bin/env python3
from __future__ import annotations

from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.formatting.rule import FormulaRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

ROOT = Path(__file__).resolve().parents[2]
DIST = ROOT / "course" / "dist"
DATE_PREFIX = "04-09-2026"
OPS_DOCX = DIST / f"{DATE_PREFIX}-So-tay-van-hanh-khoa-hoc-AI.docx"
SHARE_DOCX = DIST / f"{DATE_PREFIX}-Chuong-trinh-hoc-AI-Van-hanh-doanh-nghiep.docx"
XLSX = DIST / f"{DATE_PREFIX}-Lich-hoc-va-quan-ly-cohort-AI.xlsx"

PURPLE = "6D28D9"
DEEP = "24114F"
LAVENDER = "EDE9FE"
LIGHT = "F7F5FF"
GRAY = "5B6170"
GREEN = "DCFCE7"
YELLOW = "FEF3C7"
RED = "FEE2E2"
WHITE = "FFFFFF"
FONT = "DejaVu Sans"

WEEKS = [
    (1, "Làm chủ Claude và tư duy giao việc", "Bản đồ cơ hội AI và ba quy trình ưu tiên", "Phân loại việc phù hợp; giao việc theo mục tiêu, bối cảnh, đầu vào, đầu ra và tiêu chí; ẩn danh dữ liệu."),
    (2, "Xây trợ lý hiểu doanh nghiệp", "Claude Project, hồ sơ doanh nghiệp và checklist chất lượng", "Tổ chức nguồn; xử lý thông tin thiếu/xung đột; xây giọng thương hiệu có tiêu chí."),
    (3, "Xây hệ thống nội dung marketing", "Lịch nội dung 7 ngày và gói Facebook–email–video", "Tái sử dụng một nguồn đã xác minh; kiểm tra dữ kiện, giọng, kênh và CTA trước khi duyệt."),
    (4, "Xây trợ lý bán hàng", "Kịch bản khám phá, xử lý từ chối và chuỗi theo dõi ba chạm", "Hỗ trợ tư vấn nhưng không tự quyết giá/chiết khấu; có điều kiện dừng và bước duyệt."),
    (5, "Chăm sóc khách hàng và vận hành", "Bộ FAQ–chuyển cấp và một SOP nội bộ", "Phân loại yêu cầu; đặt ngưỡng chuyển cấp; SOP có vai trò, kiểm soát và bằng chứng hoàn tất."),
    (6, "Tự động hóa và bàn giao hệ thống", "Workflow chạy thử và hồ sơ bàn giao bộ máy AI", "Thiết kế trigger, validation, chống trùng, nhánh lỗi, log, người duyệt và runbook tắt hệ thống."),
]

SCHEDULE = [
    ("ONB", "Onboarding", date(2026, 9, 23), "20:00–20:45", "Google Meet", "Kiểm tra tài khoản, dữ liệu, quy trình đầu tiên và baseline"),
    ("W1", "Buổi chính tuần 1", date(2026, 9, 26), "19:30–22:00", "Google Meet", WEEKS[0][1]),
    ("C1", "AI Clinic tuần 1", date(2026, 9, 30), "20:00–21:00", "Google Meet", "Sửa ca thực tế tuần 1; không giảng bài mới"),
    ("W2", "Buổi chính tuần 2", date(2026, 10, 3), "19:30–22:00", "Google Meet", WEEKS[1][1]),
    ("C2", "AI Clinic tuần 2", date(2026, 10, 7), "20:00–21:00", "Google Meet", "Sửa ca thực tế tuần 2; không giảng bài mới"),
    ("W3", "Buổi chính tuần 3", date(2026, 10, 10), "19:30–22:00", "Google Meet", WEEKS[2][1]),
    ("C3", "AI Clinic tuần 3", date(2026, 10, 14), "20:00–21:00", "Google Meet", "Sửa ca thực tế tuần 3; không giảng bài mới"),
    ("W4", "Buổi chính tuần 4", date(2026, 10, 17), "19:30–22:00", "Google Meet", WEEKS[3][1]),
    ("C4", "AI Clinic tuần 4", date(2026, 10, 21), "20:00–21:00", "Google Meet", "Sửa ca thực tế tuần 4; không giảng bài mới"),
    ("W5", "Buổi chính tuần 5", date(2026, 10, 24), "19:30–22:00", "Google Meet", WEEKS[4][1]),
    ("C5", "AI Clinic tuần 5", date(2026, 10, 28), "20:00–21:00", "Google Meet", "Sửa ca thực tế tuần 5; không giảng bài mới"),
    ("W6", "Buổi chính tuần 6", date(2026, 10, 31), "19:30–22:00", "Google Meet", WEEKS[5][1]),
    ("C6", "AI Clinic tốt nghiệp", date(2026, 11, 4), "20:00–21:00", "Google Meet", "Demo, sửa hồ sơ bàn giao và chốt kế hoạch duy trì"),
]


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = FONT
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def setup_doc(doc, title, subtitle, audience):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    styles = doc.styles
    for name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DEEP)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(PURPLE)
    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(DEEP)
    styles["Heading 2"].font.bold = True
    styles["Heading 3"].font.size = Pt(10.5)
    styles["Heading 3"].font.color.rgb = RGBColor.from_string(PURPLE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("KHANH.DESIGN")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(title)
    r.bold = True; r.font.name = FONT; r.font.size = Pt(28); r.font.color.rgb = RGBColor.from_string(DEEP)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.name = FONT; r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(PURPLE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(35)
    r = p.add_run(f"{audience}\nPhiên bản 04/09/2026")
    r.font.name = FONT; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True, WHITE, 8.5)
        shade(table.rows[0].cells[i], PURPLE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, False, DEEP, 8.3)
            if len(table.rows) % 2 == 1:
                shade(cells[i], LIGHT)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_footer(doc, text):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(GRAY)


def build_ops_doc():
    doc = Document()
    setup_doc(doc, "SỔ TAY VẬN HÀNH KHÓA HỌC AI", "AI VẬN HÀNH DOANH NGHIỆP", "Tài liệu quản lý nội bộ")
    doc.add_heading("1. Mục tiêu và nguyên tắc", level=1)
    doc.add_paragraph("Vận hành một cohort tối đa 15 học viên, giúp mỗi người hoàn thành bộ máy trợ lý AI có quy trình, tiêu chí, bước duyệt và hồ sơ bàn giao. Claude là công cụ trung tâm; công cụ khác chỉ hỗ trợ khi cần.")
    add_bullets(doc, [
        "Dạy kết quả quan sát được, không dạy danh sách tính năng.",
        "Mỗi tuần phải tạo một sản phẩm nối vào bài tốt nghiệp.",
        "AI tạo bản nháp; con người chịu trách nhiệm kiểm tra và phê duyệt.",
        "Chỉ dùng dữ liệu giả lập hoặc đã được phép và ẩn danh.",
        "Không nhận mật khẩu, OTP, token, API key hoặc tài liệu khách hàng nguyên bản.",
    ])

    doc.add_heading("2. Bản đồ vận hành từ đăng ký đến tốt nghiệp", level=1)
    phases = [
        ("1", "Nhận hồ sơ", "Tự động ghi nhận riêng tư", "Hồ sơ NEW; không mời thanh toán tự động"),
        ("2", "Xét phù hợp", "Trong 24 giờ", "SẴN SÀNG / CẦN BỔ SUNG / KHÔNG PHÙ HỢP"),
        ("3", "Giữ chỗ", "Sau khi xác nhận phù hợp", "Gửi học phí, chính sách, hạn thanh toán"),
        ("4", "Onboarding", "Trước khai giảng", "Tài khoản, dữ liệu, quy trình đầu tiên, baseline"),
        ("5", "Vận hành 6 tuần", "Buổi chính + AI Clinic", "Theo dõi Đạt/Cần sửa và một hành động tiếp"),
        ("6", "Tốt nghiệp", "Sau tuần 6", "Demo, runbook, bàn giao, phản hồi và kế hoạch duy trì"),
    ]
    add_table(doc, ["Giai đoạn", "Việc", "Thời điểm", "Đầu ra kiểm soát"], phases, [1.4, 3.3, 3.5, 8.0])

    doc.add_heading("3. Quy trình xử lý học viên mới", level=1)
    add_numbered(doc, [
        "Trong 24 giờ: gọi/Zalo xác nhận bài toán, khả năng dự đủ sáu buổi, thời gian làm bài và việc dùng dữ liệu mẫu/đã ẩn danh.",
        "Đánh dấu một trạng thái: SẴN SÀNG; CẦN BỔ SUNG (ghi đúng một việc và hạn phản hồi); hoặc KHÔNG PHÙ HỢP HIỆN TẠI.",
        "Chỉ với hồ sơ SẴN SÀNG: gửi mức học phí áp dụng, chính sách hoàn phí, chi phí công cụ và hạn thanh toán.",
        "Sau xác nhận thanh toán: cấp mã HV, lịch dự kiến/chính thức, link tài liệu, kênh hỗ trợ và checklist onboarding.",
        "Không thêm học viên vào thư mục có thể xem dữ liệu của người khác; mỗi người có thư mục nộp bài riêng.",
    ])

    doc.add_heading("4. Lịch cohort dự kiến", level=1)
    doc.add_paragraph("Lịch dưới đây là phương án quản lý; chỉ chuyển trạng thái thành CHÍNH THỨC sau khi cậu chủ xác nhận và gửi học viên.")
    add_table(doc, ["Mã", "Buổi", "Ngày", "Giờ", "Nội dung"], [(c, n, d.strftime("%d/%m/%Y"), t, topic) for c,n,d,t,_,topic in SCHEDULE], [1.2, 3.2, 2.3, 2.4, 7.3])

    doc.add_heading("5. Chương trình và sản phẩm từng tuần", level=1)
    add_table(doc, ["Tuần", "Chủ đề", "Sản phẩm bắt buộc", "Năng lực chính"], [(str(w), title, product, focus) for w,title,product,focus in WEEKS], [1.0, 4.0, 5.2, 6.5])

    doc.add_heading("6. Nhịp chuẩn mỗi tuần", level=1)
    add_table(doc, ["Mốc", "Việc vận hành", "Người chịu trách nhiệm"], [
        ("Trước 24 giờ", "Kiểm tra Meet, demo, dữ liệu, rubric; gửi nhắc lịch", "Giảng viên/Trợ giảng"),
        ("Buổi chính 150'", "10' kiểm tra · 20' tư duy · 25' demo · 65' thực hành · 20' nhận xét · 10' giao bài", "Giảng viên"),
        ("Sau buổi ≤24 giờ", "Gửi recap, hạn bài, rubric và link tài nguyên", "Vận hành"),
        ("AI Clinic 60'", "5' ranh giới · 10' chọn ca · 35' xử lý ≤3 ca · 5' hành động · 5' tổng kết", "Giảng viên"),
        ("Trước tuần sau", "Chấm Đạt/Cần sửa; mỗi bài Cần sửa chỉ có một hành động trong 48 giờ", "Giảng viên/Trợ giảng"),
    ], [3.1, 9.5, 4.2])

    doc.add_heading("7. Vai trò và quyền quyết định", level=1)
    add_table(doc, ["Vai trò", "Chịu trách nhiệm", "Không làm"], [
        ("Chủ chương trình", "Duyệt học viên, lịch, học phí, hoàn phí, thay đổi phạm vi", "Không giao tự động quyết định tài chính/ngoại lệ"),
        ("Giảng viên", "Mục tiêu, demo, timebox, phản hồi theo rubric", "Không vận hành hệ thống hộ học viên"),
        ("Trợ giảng", "Theo dõi tiến độ, hỗ trợ kỹ thuật, đánh dấu xanh/vàng/đỏ", "Không nhận credential hoặc dữ liệu thô"),
        ("Học viên", "Quyền dữ liệu, kiểm tra đầu ra, quyết định sử dụng", "Không đưa dữ liệu chưa được phép vào lớp/AI"),
    ], [3.0, 8.0, 5.8])

    doc.add_heading("8. Trạng thái quản lý chuẩn", level=1)
    add_bullets(doc, [
        "Hồ sơ: NEW → CẦN BỔ SUNG / SẴN SÀNG / KHÔNG PHÙ HỢP / RÚT HỒ SƠ.",
        "Thanh toán: CHƯA GỬI → ĐÃ GỬI → ĐÃ THANH TOÁN → HOÀN PHÍ (nếu có).",
        "Bài tập: CHƯA NỘP → CẦN SỬA → ĐẠT; không dùng điểm cảm tính.",
        "Sức khỏe học viên: XANH (đúng tiến độ), VÀNG (có một vướng mắc), ĐỎ (nguy cơ bỏ dở).",
    ])

    doc.add_heading("9. Checklist trước, trong và sau buổi học", level=1)
    add_table(doc, ["Thời điểm", "Checklist tối thiểu"], [
        ("Trước 24 giờ", "Meet + link dự phòng; demo chạy lại; ảnh đầu ra dự phòng; tài liệu tuần; rubric; tắt thông báo; dữ liệu sạch"),
        ("Trước 15 phút", "Mở phòng; bật phụ đề; kiểm tra quyền bản ghi; workbook; timer; danh sách hỗ trợ"),
        ("Trong buổi", "Giữ 65 phút thực hành; phút 120 đóng phạm vi; đầu ra ở trạng thái nháp/chờ duyệt"),
        ("Sau ≤24 giờ", "Recap; hạn bài; rubric; link; cập nhật tiến độ; xóa file tạm không cần thiết"),
    ], [3.3, 13.5])

    doc.add_heading("10. Quản lý tài liệu và quyền riêng tư", level=1)
    add_bullets(doc, [
        "Drive là nơi giao đầu ra; GitHub giữ source; local chỉ là vùng dựng tạm.",
        "Tài liệu công khai/học viên không chứa bảng quản lý hồ sơ, số điện thoại hoặc ghi chú nội bộ.",
        "Tên file dùng ngày–nội dung–phiên bản; tài liệu thay thế phải giữ một bản hiện hành rõ ràng.",
        "Không công khai response sheet; chỉ người vận hành cần thiết được truy cập.",
        "Case study, ảnh, phát biểu và ghi hình cần đồng ý riêng; mặc định không dùng cho marketing.",
    ])

    doc.add_heading("11. Chỉ số theo dõi", level=1)
    add_table(doc, ["Nhóm", "Chỉ số", "Cách dùng"], [
        ("Tuyển sinh", "Hồ sơ mới, phù hợp, thanh toán", "Đo chuyển đổi; không gây khan hiếm giả"),
        ("Tham gia", "Có mặt buổi chính/Clinic", "Phát hiện người cần hỗ trợ"),
        ("Học tập", "Tỷ lệ Đạt từng tuần, tiêu chí trượt nhiều", "Sửa ví dụ/hướng dẫn, không hạ rubric"),
        ("Vận hành", "Số ca hỗ trợ, thời gian phản hồi", "Ước lượng tải cohort tiếp theo"),
        ("Kết quả", "Workflow hoàn tất, baseline trước/sau tự đo", "Không hứa năng suất khi chưa có bằng chứng"),
    ], [2.5, 6.1, 8.2])

    doc.add_heading("12. Sự cố và phương án dự phòng", level=1)
    add_bullets(doc, [
        "Claude lỗi: dùng transcript/ảnh demo, viết brief và chấm đầu ra tĩnh.",
        "Meet lỗi: chuyển link dự phòng qua kênh đã xác minh, không phát công khai.",
        "Make/n8n đổi giao diện: dạy sơ đồ trạng thái, dùng workflow JSON đã kiểm tra và không gửi thật.",
        "Lộ dữ liệu: dừng chia sẻ, không chụp/sao chép thêm, thu hồi quyền khi được phép và chuyển sang dữ liệu mẫu.",
        "Vắng trợ giảng: giảm số bài nhận xét, giữ nguyên thời gian thực hành và điểm kiểm tra an toàn.",
    ])
    add_footer(doc, "khanh.design · AI Vận Hành Doanh Nghiệp · Tài liệu nội bộ")
    DIST.mkdir(parents=True, exist_ok=True)
    doc.save(OPS_DOCX)


def build_share_doc():
    doc = Document()
    setup_doc(doc, "CHƯƠNG TRÌNH HỌC 6 TUẦN", "Xây bộ máy trợ lý AI với Claude", "Tài liệu chia sẻ học viên")
    doc.add_heading("Kết quả sau khóa học", level=1)
    doc.add_paragraph("Bạn hoàn thành một bộ máy trợ lý AI phục vụ quy trình thật của doanh nghiệp: có nguồn dữ liệu, tiêu chí chất lượng, bước người duyệt, nhánh lỗi và hồ sơ bàn giao. Khóa học không yêu cầu lập trình và không hứa thay thế con người hoặc tăng doanh thu theo tỷ lệ cố định.")
    doc.add_heading("Đối tượng phù hợp", level=1)
    add_bullets(doc, [
        "Chủ doanh nghiệp nhỏ; người làm marketing, bán hàng, chăm sóc khách hàng hoặc vận hành.",
        "Có ít nhất một công việc lặp lại và có thể dành thời gian học, thực hành hằng tuần.",
        "Sẵn sàng dùng dữ liệu mẫu hoặc dữ liệu đã được phép và ẩn danh.",
    ])
    doc.add_heading("Hình thức", level=1)
    add_bullets(doc, [
        "6 buổi trực tuyến qua Google Meet, mỗi buổi 150 phút.",
        "6 phiên AI Clinic, mỗi phiên 60 phút; sửa ca thực tế, không giảng bài mới.",
        "Học qua sản phẩm; mỗi tuần nộp một đầu ra nối vào bài tốt nghiệp.",
        "Tối đa 15 học viên.",
    ])
    doc.add_heading("Lộ trình và đầu ra", level=1)
    add_table(doc, ["Tuần", "Chủ đề", "Bạn hoàn thành"], [(str(w), title, product) for w,title,product,_ in WEEKS], [1.0, 6.0, 9.8])
    doc.add_heading("Lịch dự kiến", level=1)
    doc.add_paragraph("Lịch có thể điều chỉnh trước khi lớp xác nhận chính thức. Múi giờ: Việt Nam (UTC+7).")
    add_table(doc, ["Buổi", "Ngày", "Giờ", "Nội dung"], [(n, d.strftime("%d/%m/%Y"), t, topic) for _,n,d,t,_,topic in SCHEDULE], [4.3, 2.5, 2.5, 7.2])
    doc.add_heading("Cách học và đánh giá", level=1)
    add_bullets(doc, [
        "Mỗi sản phẩm được chấm Đạt/Cần sửa bằng tiêu chí quan sát được.",
        "Bài nộp phải có đầu vào đã làm sạch, phiên bản trước/sau, rubric tự chấm và người duyệt.",
        "Bài tốt nghiệp tích hợp sáu sản phẩm thành workflow có validation, chống trùng, nhánh lỗi, log và runbook.",
        "Nếu Cần sửa, học viên nhận một hành động cụ thể có thể hoàn tất trong 48 giờ.",
    ])
    doc.add_heading("Công cụ", level=1)
    doc.add_paragraph("Claude là công cụ AI trung tâm. Google Meet/Drive/Form hỗ trợ lớp học và tài liệu; Canva, Google Workspace, Make hoặc n8n chỉ dùng ở phần phù hợp. Học phí không bao gồm phí tài khoản hoặc công cụ bên thứ ba; mọi chi phí cần thiết phải được thông báo trước thanh toán.")
    doc.add_heading("Học phí cohort sáng lập", level=1)
    add_bullets(doc, [
        "5 suất sáng lập: 3.900.000 đồng/người.",
        "10 suất chính thức: 4.900.000 đồng/người.",
        "Thanh toán một lần sau khi hồ sơ được xác nhận phù hợp.",
        "Có thể yêu cầu hoàn phí sau buổi đầu và trước buổi thứ hai; không hoàn sau khi nhận toàn bộ template hoặc tiếp tục từ buổi hai, trừ khi chương trình bị hủy từ phía tổ chức.",
    ])
    doc.add_heading("An toàn dữ liệu", level=1)
    add_bullets(doc, [
        "Không gửi mật khẩu, OTP, token, API key hoặc tài liệu khách hàng nguyên bản.",
        "Không nhập dữ liệu nhận diện, hợp đồng, thanh toán hoặc bí mật kinh doanh khi chưa có quyền và chưa làm sạch.",
        "Mọi đầu ra đối ngoại ở trạng thái nháp/chờ người có trách nhiệm duyệt.",
        "Việc ghi hình, dùng phát biểu hoặc case study được xin phép riêng, không mặc định.",
    ])
    doc.add_heading("Đăng ký", level=1)
    doc.add_paragraph("Xem thông tin và nộp hồ sơ tại: https://khanh.design/ai-van-hanh-doanh-nghiep/\nNộp hồ sơ không bảo đảm nhận lớp. Chương trình đọc mức phù hợp trước khi mời thanh toán.")
    add_footer(doc, "khanh.design · AI Vận Hành Doanh Nghiệp · Chương trình dự kiến")
    doc.save(SHARE_DOCX)


def style_ws(ws, widths):
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for c in ws[1]:
        c.fill = PatternFill("solid", fgColor=PURPLE)
        c.font = Font(name=FONT, bold=True, color=WHITE)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 34
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    thin = Side(style="thin", color="DDD6FE")
    for row in ws.iter_rows():
        for cell in row:
            cell.font = Font(name=FONT, size=10, color=DEEP, bold=cell.row == 1)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(bottom=thin)
    for r in range(2, ws.max_row + 1):
        if r % 2 == 0:
            for c in ws[r]: c.fill = PatternFill("solid", fgColor=LIGHT)


def build_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Dashboard"
    ws.append(["CHỈ SỐ", "GIÁ TRỊ", "MỤC TIÊU/GHI CHÚ"])
    dashboard = [
        ("Tổng hồ sơ", "=COUNTA('Học viên'!B2:B16)", "Tối đa 15"),
        ("Hồ sơ sẵn sàng", '=COUNTIF(\'Học viên\'!F2:F16,"SẴN SÀNG")', "Đã xác nhận phù hợp"),
        ("Đã thanh toán", '=COUNTIF(\'Học viên\'!J2:J16,"ĐÃ THANH TOÁN")', "Chỉ sau xét phù hợp"),
        ("Đã onboarding", '=COUNTIF(\'Học viên\'!L2:L16,"HOÀN TẤT")', "Trước 23/09/2026"),
        ("Học viên XANH", '=COUNTIF(\'Tiến độ\'!J2:J16,"XANH")', "Đúng tiến độ"),
        ("Học viên VÀNG", '=COUNTIF(\'Tiến độ\'!J2:J16,"VÀNG")', "Có một vướng mắc cần xử lý"),
        ("Học viên ĐỎ", '=COUNTIF(\'Tiến độ\'!J2:J16,"ĐỎ")', "Nguy cơ bỏ dở"),
        ("Doanh thu đã nhận", "=SUM('Học viên'!K2:K16)", "VNĐ"),
        ("Chỗ còn lại", "=15-COUNTA('Học viên'!B2:B16)", "Không dùng để tạo khan hiếm giả"),
    ]
    for row in dashboard: ws.append(row)
    style_ws(ws, [27, 20, 48])
    ws["B9"].number_format = '#,##0 "đ"'
    ws["A12"] = "LỊCH ĐANG Ở TRẠNG THÁI DỰ KIẾN — chỉ công bố khi chủ chương trình xác nhận."
    ws.merge_cells("A12:C12")
    ws["A12"].fill = PatternFill("solid", fgColor=YELLOW)
    ws["A12"].font = Font(name=FONT, bold=True, color=DEEP)

    ws = wb.create_sheet("Học viên")
    ws.append(["Mã HV", "Họ tên", "Điện thoại", "Email", "Ngành", "Trạng thái hồ sơ", "Ngày liên hệ", "Hành động tiếp", "Gói học phí", "Thanh toán", "Đã nhận (VNĐ)", "Onboarding", "Ghi chú tối thiểu"])
    for i in range(1, 16): ws.append([f"HV{i:02d}", "", "", "", "", "NEW", "", "Liên hệ trong 24 giờ", "", "CHƯA GỬI", 0, "CHƯA", ""])
    style_ws(ws, [10,22,17,27,22,23,15,30,20,20,18,16,34])
    for r in range(2, 17):
        ws.cell(r,7).number_format = "dd/mm/yyyy"
        ws.cell(r,11).number_format = '#,##0 "đ"'
    for formula in ["NEW,CẦN BỔ SUNG,SẴN SÀNG,KHÔNG PHÙ HỢP,RÚT HỒ SƠ", "CHƯA GỬI,ĐÃ GỬI,ĐÃ THANH TOÁN,HOÀN PHÍ", "CHƯA,ĐANG LÀM,HOÀN TẤT"]:
        pass
    dv1=DataValidation(type="list",formula1='"NEW,CẦN BỔ SUNG,SẴN SÀNG,KHÔNG PHÙ HỢP,RÚT HỒ SƠ"'); ws.add_data_validation(dv1); dv1.add("F2:F16")
    dv2=DataValidation(type="list",formula1='"CHƯA GỬI,ĐÃ GỬI,ĐÃ THANH TOÁN,HOÀN PHÍ"'); ws.add_data_validation(dv2); dv2.add("J2:J16")
    dv3=DataValidation(type="list",formula1='"CHƯA,ĐANG LÀM,HOÀN TẤT"'); ws.add_data_validation(dv3); dv3.add("L2:L16")
    dv4=DataValidation(type="list",formula1='"Sáng lập 3.900.000,Chính thức 4.900.000"'); ws.add_data_validation(dv4); dv4.add("I2:I16")

    ws = wb.create_sheet("Lịch học")
    ws.append(["Mã", "Loại buổi", "Ngày", "Thứ", "Giờ", "Thời lượng", "Hình thức", "Chủ đề/Kết quả", "Trạng thái", "Link Meet", "Link tài liệu", "Việc trước 24h"])
    weekdays = {0:"Thứ Hai",1:"Thứ Ba",2:"Thứ Tư",3:"Thứ Năm",4:"Thứ Sáu",5:"Thứ Bảy",6:"Chủ Nhật"}
    for code,name,d,t,channel,topic in SCHEDULE:
        duration = 45 if code=="ONB" else (60 if code.startswith("C") else 150)
        ws.append([code,name,d,weekdays[d.weekday()],t,duration,channel,topic,"DỰ KIẾN","","","Kiểm tra link, demo, dữ liệu, rubric và nhắc lịch"])
    style_ws(ws, [9,23,14,13,16,12,17,43,16,30,30,43])
    for r in range(2, ws.max_row+1): ws.cell(r,3).number_format="dd/mm/yyyy"
    dv=DataValidation(type="list",formula1='"DỰ KIẾN,ĐÃ XÁC NHẬN,HOÀN TẤT,HOÃN"'); ws.add_data_validation(dv); dv.add(f"I2:I{ws.max_row}")

    ws = wb.create_sheet("Tiến độ")
    ws.append(["Mã HV", "Họ tên", "Onboarding", "Tuần 1", "Tuần 2", "Tuần 3", "Tuần 4", "Tuần 5", "Tuần 6", "Sức khỏe", "Hành động tiếp", "Hạn", "Ghi chú"])
    for i in range(1,16): ws.append([f"HV{i:02d}", f"='Học viên'!B{i+1}", "CHƯA", "CHƯA NỘP", "CHƯA NỘP", "CHƯA NỘP", "CHƯA NỘP", "CHƯA NỘP", "CHƯA NỘP", "", "", "", ""])
    style_ws(ws, [10,22,16,16,16,16,16,16,16,13,32,14,34])
    dv=DataValidation(type="list",formula1='"CHƯA NỘP,CẦN SỬA,ĐẠT"'); ws.add_data_validation(dv); dv.add("D2:I16")
    dv2=DataValidation(type="list",formula1='"CHƯA,ĐANG LÀM,HOÀN TẤT"'); ws.add_data_validation(dv2); dv2.add("C2:C16")
    dv3=DataValidation(type="list",formula1='"XANH,VÀNG,ĐỎ"'); ws.add_data_validation(dv3); dv3.add("J2:J16")
    ws.conditional_formatting.add("J2:J16", FormulaRule(formula=['J2="XANH"'],fill=PatternFill("solid",fgColor=GREEN)))
    ws.conditional_formatting.add("J2:J16", FormulaRule(formula=['J2="VÀNG"'],fill=PatternFill("solid",fgColor=YELLOW)))
    ws.conditional_formatting.add("J2:J16", FormulaRule(formula=['J2="ĐỎ"'],fill=PatternFill("solid",fgColor=RED)))

    ws = wb.create_sheet("Công việc vận hành")
    ws.append(["Mã", "Giai đoạn", "Công việc", "Hạn/mốc", "Người phụ trách", "Trạng thái", "Bằng chứng/Link", "Ghi chú"])
    tasks = [
        ("OP01","Tuyển sinh","Theo dõi website, form và hồ sơ mới","Liên tục","Trợ lý","ĐANG LÀM","","Chỉ báo khi có hồ sơ/lỗi thật"),
        ("OP02","Tuyển sinh","Liên hệ hồ sơ mới","≤24 giờ từ đăng ký","Chủ chương trình","CHƯA","","Xác nhận phù hợp trước thanh toán"),
        ("OP03","Chuẩn bị","Chốt lịch cohort và gửi học viên","Trước 20/09","Chủ chương trình","CHƯA","","Đổi lịch từ DỰ KIẾN sang ĐÃ XÁC NHẬN"),
        ("OP04","Chuẩn bị","Tạo Meet chính + dự phòng","Trước 22/09","Vận hành","CHƯA","","Không phát link công khai"),
        ("OP05","Chuẩn bị","Tạo thư mục riêng cho từng học viên","Trước onboarding","Vận hành","CHƯA","","Không cho xem chéo"),
        ("OP06","Onboarding","Chạy checklist 45 phút","23/09 20:00","Giảng viên","CHƯA","","Tài khoản, dữ liệu, baseline"),
        ("OP07","Hằng tuần","Kiểm tra demo, dữ liệu, rubric và link","Trước buổi 24h","Giảng viên/Trợ giảng","CHƯA","","Có phương án dự phòng"),
        ("OP08","Hằng tuần","Gửi recap và hạn bài","Sau buổi ≤24h","Vận hành","CHƯA","","Không chứa dữ liệu riêng"),
        ("OP09","Hằng tuần","Chấm Đạt/Cần sửa","Trước buổi kế","Giảng viên","CHƯA","","Một hành động sửa trong 48h"),
        ("OP10","Tốt nghiệp","Demo, runbook và bàn giao","04/11","Giảng viên","CHƯA","","Có chủ sở hữu và lịch rà soát"),
        ("OP11","Sau khóa","Phản hồi, chỉ số và quyền case study","Sau khóa 7 ngày","Chủ chương trình","CHƯA","","Case study cần đồng ý riêng"),
    ]
    for row in tasks: ws.append(row)
    style_ws(ws, [9,17,42,20,23,17,32,38])
    dv=DataValidation(type="list",formula1='"CHƯA,ĐANG LÀM,CHỜ,HOÀN TẤT,HOÃN"'); ws.add_data_validation(dv); dv.add(f"F2:F{ws.max_row}")

    ws = wb.create_sheet("Chương trình 6 tuần")
    ws.append(["Tuần", "Chủ đề", "Sản phẩm bắt buộc", "Năng lực/Phạm vi", "Trạng thái tài liệu"])
    for w,title,product,focus in WEEKS: ws.append([w,title,product,focus,"SẴN SÀNG"])
    style_ws(ws, [10,31,48,65,20])

    ws = wb.create_sheet("Chỉ mục tài liệu")
    ws.append(["Nhóm", "Tài liệu", "Đối tượng", "Định dạng", "Đường dẫn trong repo", "Trạng thái", "Quy tắc chia sẻ"])
    docs = [
        ("Quản trị","Sổ tay vận hành khóa học AI","Nội bộ","DOCX/PDF",f"course/dist/{OPS_DOCX.name}","HIỆN HÀNH","Không gửi bảng PII cho học viên"),
        ("Chia sẻ","Chương trình học 6 tuần","Học viên/Ứng viên","PDF",f"course/dist/{SHARE_DOCX.with_suffix('.pdf').name}","HIỆN HÀNH","Có thể chia sẻ"),
        ("Lịch","Lịch học và quản lý cohort","Nội bộ","XLSX",f"course/dist/{XLSX.name}","HIỆN HÀNH","Giới hạn quyền truy cập"),
        ("Giảng viên","Giáo án chi tiết Horus 6 tuần","Giảng viên","DOCX/PDF","course/dist/04-09-2026-Giao-an-chi-tiet-Horus-AI-Van-hanh-doanh-nghiep.pdf","HIỆN HÀNH","Bản PDF để dạy; DOCX để chỉnh sửa"),
        ("Học viên","Workbook","Học viên","PDF","course/dist/04-09-2026-Workbook-AI-Van-hanh-doanh-nghiep.pdf","HIỆN HÀNH","Chia sẻ sau xác nhận"),
        ("Workshop","Slide workshop","Người tham dự","PPTX/PDF","course/dist/04-09-2026-Workshop-Xay-tro-ly-AI-dau-tien-voi-Claude.pdf","HIỆN HÀNH","Có thể chia sẻ theo chương trình"),
        ("Onboarding","Checklist intake","Vận hành","Source/PDF gộp","course/onboarding/intake-checklist.md","HIỆN HÀNH","Không public response sheet"),
        ("Website","Landing page","Công khai","HTML","ai-van-hanh-doanh-nghiep/","LIVE","https://khanh.design/ai-van-hanh-doanh-nghiep/"),
    ]
    for row in docs: ws.append(row)
    style_ws(ws, [17,38,22,16,67,18,42])

    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    wb.calculation.calcMode = "auto"
    DIST.mkdir(parents=True, exist_ok=True)
    wb.save(XLSX)


def main():
    build_ops_doc(); build_share_doc(); build_xlsx()
    for path in [OPS_DOCX, SHARE_DOCX, XLSX]: print(path)


if __name__ == "__main__": main()
